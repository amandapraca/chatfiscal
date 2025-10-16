import streamlit as st
import pandas as pd
from agent_manager import AgentManager
from llm_utils import gerar_resposta_llm
from main.dicas_corujito import gerar_dica_corujito
from main.interface import (
    montar_interface,
    exibir_resposta_agente,
    exibir_rodape,
    mostrar_alerta,
    mostrar_erro,
    mostrar_sucesso,
    boas_vindas,
    introducao_chatfiscal,
    exibir_dica_corujito
)

# ConfiguraÃ§Ã£o da pÃ¡gina
st.set_page_config(page_title="ChatFiscal", layout="wide")
montar_interface()

# InicializaÃ§Ã£o do agente pai
manager = AgentManager()

# InicializaÃ§Ã£o da sessÃ£o
if "df" not in st.session_state:
    st.session_state["df"] = None
if "historico" not in st.session_state:
    st.session_state["historico"] = []
if "arquivo_carregado" not in st.session_state:
    st.session_state["arquivo_carregado"] = False
if "dica_exibida" not in st.session_state:
    st.session_state["dica_exibida"] = False

# CriaÃ§Ã£o das abas
abas = st.tabs([
    "ğŸ“Š Dados & Perguntas",
    "ğŸ“š HistÃ³rico",
    "ğŸ” Auditoria",
    "ğŸ“ˆ VisualizaÃ§Ãµes",
    "ğŸ§  Painel Inteligente",
])

# ğŸ§© Aba 1 â€” Upload de Arquivos e Perguntas ao Agente
with abas[0]:
    st.subheader("ğŸ“Š Dados & Perguntas")
    st.markdown("Envie um ou mais arquivos fiscais para anÃ¡lise. Aceitamos arquivos CSV ou XML.")

    arquivos = st.file_uploader(
        "ğŸ“ Escolha os arquivos",
        type=["csv", "xml"],
        accept_multiple_files=True,
        key="upload_arquivos_fiscais",
    )

    dfs = []
    for arquivo in arquivos:
        df_individual = manager.carregar_arquivo(arquivo)
        if isinstance(df_individual, pd.DataFrame) and not df_individual.empty:
            dfs.append(df_individual)

    if dfs:
        df = pd.concat(dfs, ignore_index=True)
        st.session_state["df"] = df
        st.session_state["arquivo_carregado"] = True
        st.session_state["dica_exibida"] = False

        mostrar_sucesso(f"{len(dfs)} arquivo(s) carregado(s) com sucesso!")
        st.dataframe(df)

        if not st.session_state["dica_exibida"]:
            dica = gerar_dica_corujito(df)
            exibir_dica_corujito(dica)
            st.session_state["dica_exibida"] = True

        pergunta = st.text_input("Digite uma pergunta fiscal para o agente:", key="campo_pergunta_agente")

        if pergunta and st.button("Enviar", key="botao_enviar_pergunta"):
            with st.spinner("Analisando sua pergunta..."):
                resposta = gerar_resposta_llm(pergunta, df)
                st.session_state["historico"].append((pergunta, resposta))
                exibir_resposta_agente(pergunta, resposta)
    else:
        mostrar_alerta("Nenhum arquivo carregado ainda. VocÃª pode simular dados ou enviar arquivos reais.")

# ğŸ§© Aba 2 â€” HistÃ³rico de Perguntas e Respostas
with abas[1]:
    st.subheader("ğŸ“š HistÃ³rico de Perguntas")

    if st.session_state["historico"]:
        for i, (pergunta, resposta) in enumerate(st.session_state["historico"], 1):
            st.markdown(f"**{i}. Pergunta:** {pergunta}")
            st.markdown(f"**Resposta:** {resposta}")
            st.markdown("---")

        if st.button("ğŸ§¹ Limpar histÃ³rico", key="botao_limpar_historico"):
            st.session_state["historico"] = []
            mostrar_sucesso("HistÃ³rico limpo com sucesso!")

        try:
            from docx import Document
            from io import BytesIO

            doc = Document()
            doc.add_heading("RelatÃ³rio de Perguntas e Respostas", level=1)

            for i, (p, r) in enumerate(st.session_state["historico"], 1):
                doc.add_paragraph(f"{i}. Pergunta: {p}")
                doc.add_paragraph(f"   Resposta: {r}")
                doc.add_paragraph("")

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="ğŸ“¤ Exportar histÃ³rico como relatÃ³rio (.docx)",
                data=buffer,
                file_name="relatorio_chatfiscal.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="botao_exportar_docx",
            )
        except Exception as e:
            mostrar_erro(f"Erro ao gerar relatÃ³rio: {e}")
    else:
        mostrar_alerta("Nenhuma pergunta registrada ainda. FaÃ§a uma anÃ¡lise para comeÃ§ar.")

# ğŸ§© Aba 3 â€” Auditoria
with abas[2]:
    st.header("ğŸ” Auditoria")
    if st.button("Validar Arquivo"):
        resultado = manager.validar_arquivo()
        st.text(resultado)

# ğŸ§© Aba 4 â€” VisualizaÃ§Ãµes
with abas[3]:
    st.header("ğŸ“ˆ VisualizaÃ§Ãµes")
    pergunta = st.text_input("FaÃ§a uma pergunta sobre os dados carregados")
    if st.button("Gerar Resposta"):
        resposta = manager.gerar_resposta(pergunta)
        st.text(resposta)

# ğŸ§© Aba 5 â€” Painel Inteligente (placeholder)
with abas[4]:
    st.header("ğŸ§  Painel Inteligente")
    st.info("Em breve: KPIs fiscais, grÃ¡ficos e insights inteligentes.")

# ğŸ§© RodapÃ© institucional
exibir_rodape()
